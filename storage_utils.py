# storage_utils.py
from __future__ import annotations

import os, json, io, traceback, re
from io import BytesIO
from typing import Optional, Tuple, List
from datetime import datetime, timedelta, timezone

from azure.storage.blob import (
    BlobServiceClient,
    generate_blob_sas,
    BlobSasPermissions,
    ContentSettings,
)
from azure.core.exceptions import (
    ResourceNotFoundError,
    HttpResponseError,
    ClientAuthenticationError,
)

# PDF 관련 (옵션: 지금은 DOCX 저장이 메인)
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
except Exception:
    Document = None
    Pt = None
    qn = None

# ── ENV ─────────────────────────────────────────────────────────
CONN_STR = (os.getenv("AZURE_STORAGE_CONNECTION_STRING") or "").strip()
ACCOUNT_NAME = (os.getenv("AZURE_STORAGE_ACCOUNT") or "").strip()
ACCOUNT_KEY = (os.getenv("AZURE_STORAGE_KEY") or "").strip()
CONTAINER = (os.getenv("AZURE_STORAGE_CONTAINER") or "news").strip()

_client: Optional[BlobServiceClient] = None


# ── Core Client ─────────────────────────────────────────────────
def _svc() -> BlobServiceClient:
    """
    우선순위:
      1) Connection String
      2) AccountName + AccountKey
    """
    global _client
    if _client:
        return _client

    if CONN_STR:
        _client = BlobServiceClient.from_connection_string(CONN_STR)
        return _client

    if not ACCOUNT_NAME or not ACCOUNT_KEY:
        raise RuntimeError(
            "스토리지 인증 정보가 없습니다. "
            "AZURE_STORAGE_CONNECTION_STRING 또는 (AZURE_STORAGE_ACCOUNT + AZURE_STORAGE_KEY) 를 설정하세요."
        )

    account_url = f"https://{ACCOUNT_NAME}.blob.core.windows.net"
    _client = BlobServiceClient(account_url=account_url, credential=ACCOUNT_KEY)
    return _client


# ── 폰트 헬퍼 (PDF용 — 지금은 DOCX가 메인) ─────────────────────────
def _download_noto_font(font_dir: str = "./fonts") -> Optional[str]:
    font_url = "https://github.com/googlefonts/noto-cjk/raw/main/Sans/OTF/KR/NotoSansKR-Regular.otf"
    os.makedirs(font_dir, exist_ok=True)
    font_path = os.path.join(font_dir, "NotoSansKR-Regular.otf")
    if not os.path.exists(font_path):
        try:
            import requests

            r = requests.get(font_url)
            r.raise_for_status()
            with open(font_path, "wb") as f:
                f.write(r.content)
        except Exception as e:
            print(f"폰트 다운로드 실패: {e}")
            return None
    return font_path


def _find_korean_font_path() -> Optional[str]:
    env_path = (os.getenv("KOREAN_FONT_PATH") or "").strip()
    if env_path and os.path.exists(env_path):
        return env_path
    noto_font = _download_noto_font()
    if noto_font and os.path.exists(noto_font):
        return noto_font
    candidates = [
        "/System/Library/Fonts/AppleSDGothicNeo.ttc",
        "/Library/Fonts/Apple SD Gothic Neo.ttf",
        "/Library/Fonts/NotoSansCJKkr-Regular.otf",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJKkr-Regular.otf",
        "C:\\Windows\\Fonts\\malgun.ttf",
        "C:\\Windows\\Fonts\\malgunbd.ttf",
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansKR-Regular.otf",
    ]
    for p in candidates:
        if os.path.exists(p):
            return p
    return None


def _register_korean_font(font_name: str = "KFONT") -> Optional[str]:
    try:
        fp = _find_korean_font_path()
        if not fp:
            return None
        pdfmetrics.registerFont(TTFont(font_name, fp))
        return font_name
    except Exception:
        return None


# ── PDF 생성 (옵션) ─────────────────────────────────────────────
def generate_pdf_bytes(
    items: List[dict],
    *,
    title: str,
    query: str,
    freshness: str,
    market: str,
    generated_at: Optional[str] = None,
) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36
    )

    font_name = _register_korean_font() or "Helvetica"
    styles = getSampleStyleSheet()

    base = {"fontName": font_name, "fontSize": 11, "leading": 16}
    base_title = {
        k: v for k, v in base.items() if k not in ("fontSize", "leading", "spaceAfter")
    }
    style_title = ParagraphStyle(
        "title", **base_title, fontSize=16, leading=22, spaceAfter=12
    )
    style_sub = ParagraphStyle(
        "sub", **base, textColor=colors.HexColor("#555555"), spaceAfter=8
    )
    style_cell = ParagraphStyle("cell", **base)

    story: List = []
    story.append(Paragraph(title, style_title))
    subtitle = f"Query: {query}  |  Freshness: {freshness}  |  Market: {market}"
    if generated_at:
        subtitle += f"  |  Generated: {generated_at}"
    story.append(Paragraph(subtitle, style_sub))
    story.append(Spacer(1, 6))

    data = [["제목", "매체/시간", "링크"]]
    for it in items:
        title_p = Paragraph((it.get("title") or "").replace("\n", " "), style_cell)
        meta = f"{it.get('source') or '-'} / {it.get('published') or '-'}"
        meta_p = Paragraph(meta, style_cell)
        url = it.get("url") or "-"
        url_p = Paragraph(url.replace(" ", ""), style_cell)
        data.append([title_p, meta_p, url_p])

    table = Table(data, colWidths=[250, 120, 150])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f3f6")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#222222")),
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#cccccc")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                (
                    "ROWBACKGROUNDS",
                    (0, 1),
                    (-1, -1),
                    [colors.white, colors.HexColor("#fbfcfe")],
                ),
            ]
        )
    )
    story.append(table)

    if font_name == "Helvetica":
        story.append(Spacer(1, 6))
        warn = "⚠️ 한글 폰트가 없어 일부 문자가 깨질 수 있습니다. .env의 KOREAN_FONT_PATH에 한글 폰트 경로를 지정하세요."
        story.append(Paragraph(warn, style_sub))

    doc.build(story)
    return buf.getvalue()


# ── DOCX 생성 ──────────────────────────────────────────────────
def generate_docx_bytes(
    items: List[dict],
    *,
    title: str = "뉴스 스크랩 리포트",
    query: str = "",
    freshness: str = "",
    market: str = "",
    generated_at: Optional[str] = None,
) -> bytes:
    """
    뉴스 항목을 docx(워드) 파일로 변환하여 바이트 반환. (한글 호환)
    - 인자 대부분 선택형으로 만들어 기존/새 호출 모두 수용
    """
    if Document is None:
        raise RuntimeError(
            "python-docx가 설치되어 있지 않습니다. 'python-docx' 패키지를 설치하세요."
        )

    doc = Document()
    # 한글 기본 폰트 시도
    try:
        style = doc.styles["Normal"]
        font = style.font
        font.name = "맑은 고딕"
        font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    except Exception:
        pass

    doc.add_heading(title, level=0)
    subtitle = (
        f"Query: {query}"
        + (f"  |  Freshness: {freshness}" if freshness else "")
        + (f"  |  Market: {market}" if market else "")
    )
    if generated_at:
        subtitle += f"  |  Generated: {generated_at}"
    doc.add_paragraph(subtitle.strip("  |"))

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "제목"
    hdr_cells[1].text = "매체/시간"
    hdr_cells[2].text = "링크"

    for it in items:
        row_cells = table.add_row().cells
        row_cells[0].text = (it.get("title") or "").replace("\n", " ")
        meta = f"{it.get('source') or '-'} / {it.get('published') or '-'}"
        row_cells[1].text = meta
        row_cells[2].text = it.get("url") or "-"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── 파일명 유틸 (옵션) ──────────────────────────────────────────
def make_pdf_filename_from_query(query: str, include_date: bool = True) -> str:
    q = (query or "").strip() or "pressm"
    slug = re.sub(r"[^0-9A-Za-z가-힣]+", "_", q).strip("_")
    if include_date:
        today = datetime.now().strftime("%Y%m%d")
        return f"pdf/{slug}_{today}.pdf"
    return f"pdf/{slug}.pdf"


def make_docx_filename_from_query(query: str, include_date: bool = True) -> str:
    q = (query or "").strip() or "pressm"
    slug = re.sub(r"[^0-9A-Za-z가-힣]+", "_", q).strip("_")
    if include_date:
        today = datetime.now().strftime("%Y%m%d")
        return f"docx/{slug}_{today}.docx"
    return f"docx/{slug}.docx"


# ── Container ───────────────────────────────────────────────────
def ensure_container():
    svc = _svc()
    if not CONTAINER.islower():
        raise RuntimeError(
            "AZURE_STORAGE_CONTAINER는 반드시 소문자여야 합니다. (예: news)"
        )
    try:
        svc.create_container(CONTAINER)
    except Exception as e:
        s = str(e)
        if "ContainerAlreadyExists" in s or "Conflict" in s:
            return
        raise


def blob_exists(container: str, blob_path: str) -> bool:
    return _svc().get_blob_client(container, blob_path).exists()


# ── Upload / Download ───────────────────────────────────────────
def upload_json(obj, *, prefix: str = "news/json") -> Tuple[str, str]:
    ensure_container()
    now = datetime.now().strftime("%Y-%m-%d")
    ts = datetime.now().strftime("%H%M%S")
    blob_path = f"{prefix}/{now}/pressm_{ts}.json"

    data = json.dumps(obj, ensure_ascii=False, indent=2).encode("utf-8")
    content = ContentSettings(content_type="application/json; charset=utf-8")

    bc = _svc().get_blob_client(CONTAINER, blob_path)
    try:
        bc.upload_blob(data, overwrite=True, content_settings=content)
        return CONTAINER, blob_path
    except ClientAuthenticationError as e:
        raise RuntimeError(f"[auth] 인증 실패: {e}")
    except HttpResponseError as e:
        raise RuntimeError(f"[http] 업로드 실패: {e}")
    except Exception as e:
        raise RuntimeError(f"[blob] 업로드 실패: {e}\n{traceback.format_exc()}")


def download_bytes(container: str, blob: str) -> bytes:
    bc = _svc().get_blob_client(container=container, blob=blob)
    try:
        stream = bc.download_blob()
        return stream.readall()
    except ResourceNotFoundError as e:
        raise RuntimeError(f"Blob을 찾을 수 없습니다: {container}/{blob} — {e}")
    except Exception as e:
        raise RuntimeError(f"Blob 다운로드 실패: {e}")


# ── URL / SAS ───────────────────────────────────────────────────
def public_blob_url(container: str, blob: str) -> str:
    acct = ACCOUNT_NAME or (
        _svc().account_name if hasattr(_svc(), "account_name") else ""
    )
    return f"https://{acct}.blob.core.windows.net/{container}/{blob}"


def sas_url(container: str, blob: str, minutes: int = 120) -> Optional[str]:
    acct = ACCOUNT_NAME or (
        _svc().account_name if hasattr(_svc(), "account_name") else ""
    )
    if not acct:
        return None
    ak = ACCOUNT_KEY
    if not ak and CONN_STR:
        # 연결문자열만 있고 ACCOUNT_KEY는 없는 경우 → SAS 생략(불가)
        return None
    if not ak:
        return None

    token = generate_blob_sas(
        account_name=acct,
        account_key=ak,
        container_name=container,
        blob_name=blob,
        permission=BlobSasPermissions(read=True),
        expiry=datetime.utcnow() + timedelta(minutes=minutes),
    )
    return f"https://{acct}.blob.core.windows.net/{container}/{blob}?{token}"


# ── 버전 파일명 유틸 ────────────────────────────────────────────
def _next_version_name(container: str, base_path: str) -> str:
    """
    base_path: 'news_pdf/2025-10-30.docx' 가 이미 있으면
    'news_pdf/2025-10-30 (1).docx', '... (2).docx' ... 로 증가
    """
    if not blob_exists(container, base_path):
        return base_path

    if "." in base_path:
        stem, ext = base_path.rsplit(".", 1)
        ext = "." + ext
    else:
        stem, ext = base_path, ""

    n = 1
    while True:
        candidate = f"{stem} ({n}){ext}"
        if not blob_exists(container, candidate):
            return candidate
        n += 1


# ── DOCX 업로드 (요구사항 핵심) ──────────────────────────────────
def upload_docx_report(
    items: List[dict], *, query: str = "", kst_date: datetime | None = None
) -> tuple[str, str, str]:
    """
    DOCX 리포트를 news_pdf/YYYY-MM-DD.docx 로 업로드.
    동일 날짜 파일이 있으면 (1), (2) 등 버전 부여.
    return: (container, blob_path, sas_link)
    """
    ensure_container()

    KST = timezone(timedelta(hours=9))
    d = (kst_date or datetime.now(KST)).date().isoformat()  # YYYY-MM-DD
    base_blob = f"news_pdf/{d}.docx"
    blob_path = _next_version_name(CONTAINER, base_blob)

    docx_bytes = generate_docx_bytes(
        items,
        title=f"뉴스 스크랩 리포트 — {d}",
        query=query,
        freshness="",
        market="",
        generated_at=datetime.now(KST).strftime("%Y-%m-%d %H:%M:%S KST"),
    )

    content = ContentSettings(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    _svc().get_blob_client(CONTAINER, blob_path).upload_blob(
        docx_bytes, overwrite=False, content_settings=content
    )

    link = sas_url(CONTAINER, blob_path, minutes=120) or public_blob_url(
        CONTAINER, blob_path
    )
    return CONTAINER, blob_path, link


# ── 임의 DOCX 업로드 (이름 지정) ─────────────────────────────────
def upload_docx(docx_bytes: bytes, blob_path: str) -> Tuple[str, str]:
    ensure_container()
    content = ContentSettings(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    bc = _svc().get_blob_client(CONTAINER, blob_path)
    try:
        bc.upload_blob(docx_bytes, overwrite=True, content_settings=content)
        return CONTAINER, blob_path
    except ClientAuthenticationError as e:
        raise RuntimeError(f"[auth] 인증 실패: {e}")
    except HttpResponseError as e:
        raise RuntimeError(f"[http] 업로드 실패: {e}")
    except Exception as e:
        raise RuntimeError(f"[blob] 업로드 실패: {e}\n{traceback.format_exc()}")


# ── CSV 변환 ─────────────────────────────────────────────────────
def to_csv_bytes(items: List[dict]) -> bytes:
    import csv

    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["title", "source", "published", "url", "snippet"])
    for it in items:
        w.writerow(
            [
                it.get("title", ""),
                it.get("source", ""),
                it.get("published", ""),
                it.get("url", ""),
                it.get("snippet", ""),
            ]
        )
    return buf.getvalue().encode("utf-8")
